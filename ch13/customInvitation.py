#! python3

import docx, os, argparse
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sys import exit

DEFAULT_STYLE = {'alignment': WD_ALIGN_PARAGRAPH.CENTER, 'font_name': 'Lucida Calligraphy', 'font_size': 12}

def apply_style(para, style):
	para.alignment = style['alignment']
	font = para.runs[0].font
	font.name = style['font_name']
	font.size = docx.shared.Pt(style['font_size'])

def apply_default_style(para):
	apply_style(para, DEFAULT_STYLE)
	pass

if __name__ == '__main__':
	parser = argparse.ArgumentParser()
	parser.add_argument('guests')
	args = parser.parse_args()

	if not os.path.exists(args.guests):
		print('Unable to find {}. Exiting the program..'.format(args.guests))
		exit()
	
	doc = docx.Document()
	with open(args.guests) as guestList: 
		allGuests = guestList.readlines()
		for guest in allGuests:
			intro = doc.add_paragraph('It would be a pleasure to have the company of')
			apply_default_style(intro)
			
			nameStyle = {'alignment': WD_ALIGN_PARAGRAPH.CENTER, 'font_name': 'cooper black', 'font_size': 26}
			name = doc.add_paragraph(guest.strip())
			apply_style(name, nameStyle)
			
			address = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
			apply_default_style(address)
			
			dateStyle = {'alignment': WD_ALIGN_PARAGRAPH.CENTER, 'font_name': 'Georgia', 'font_size': 16}
			date = doc.add_paragraph('April 1st')
			apply_style(date, dateStyle)
			
			pageBr = date.add_run()
			pageBr.add_break(docx.enum.text.WD_BREAK.PAGE)
	
	outputFile = 'custom_invitations.docx'
	doc.save(outputFile)
	print('custom invitations saved to {}'.format(outputFile))